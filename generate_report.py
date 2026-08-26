#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
สคริปต์สร้างรายงาน "สรุปการตรวจ รอยสัก-เจาะหู" อัตโนมัติ

วิธีใช้ (ทำทุกรอบที่ต้องออกรายงานใหม่):
  1. แก้ไฟล์ config.json ให้ตรงกับรอบที่ตรวจ (ครั้งที่ / ห้อง / สรุปผล / รายชื่อ)
  2. ใส่รูปในโฟลเดอร์ photos/ชาย/ และ photos/หญิง/ แบบแบนราบ (ไม่ต้องแยกโฟลเดอร์ย่อย)
     สคริปต์เรียงตามชื่อไฟล์ (natural sort) แล้วหั่นกลุ่มตาม group_size ใน config.json
  3. รันคำสั่ง:  python3 generate_report.py
  4. ได้ไฟล์ Word ในโฟลเดอร์ output/ พร้อมพิมพ์/เซ็นชื่อ
"""

import io
import json
import re
import sys
from pathlib import Path

from PIL import Image as PilImage
from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent
CONFIG_PATH = BASE_DIR / "config.json"

FONT_NAME = "TH Sarabun New"
BASE_SIZE = 16

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}


def image_to_stream(img_path: Path) -> io.BytesIO:
    """ส่งรูปผ่าน Pillow เพื่อ strip EXIF ที่มีค่าผิดปกติ ก่อนส่งให้ python-docx"""
    buf = io.BytesIO()
    with PilImage.open(img_path) as pil_img:
        pil_img.convert("RGB").save(buf, format="JPEG", quality=95)
    buf.seek(0)
    return buf


def natural_sort_key(path: Path):
    """เรียงไฟล์โดยจัดการเลขในชื่อถูกต้อง เช่น img2 < img10"""
    parts = re.split(r"(\d+)", path.name.lower())
    return [int(p) if p.isdigit() else p for p in parts]


# ---------- ตัวช่วยเรื่องฟอนต์ไทย ----------

def set_thai_font(run, name=FONT_NAME, size=BASE_SIZE, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def add_text(paragraph, text, size=BASE_SIZE, bold=False):
    run = paragraph.add_run(text)
    set_thai_font(run, size=size, bold=bold)
    return run


def checkbox(checked):
    return "☑" if checked else "☐"


def get_available_width_inches(doc):
    section = doc.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    return width / 914400  # EMU -> inch


def add_image_row(paragraph, image_paths, available_width_in):
    # ล็อคความกว้างแบบ 2 รูปต่อแถวเสมอ ทำให้ทุกรูปมีขนาดเท่ากันแม้แถวสุดท้ายมีรูปเดียว
    each_w = min((available_width_in / 2) * 0.90, 3.0)
    for i, img_path in enumerate(image_paths):
        if i > 0:
            add_text(paragraph, " ")
        run = paragraph.add_run()
        run.add_picture(image_to_stream(img_path), width=Inches(each_w))


# ---------- สร้างหน้าสรุป (หน้าแรก) ----------

def build_summary_page(doc, cfg):
    logo_path = cfg.get("logo_path")
    if logo_path and (BASE_DIR / logo_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(BASE_DIR / logo_path), width=Inches(2.2))

    header = cfg["header"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, f"ครั้งที่ {header['ครั้งที่']} ภาคเรียนที่ {header['ภาคเรียน']}", size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, f"สรุปการตรวจ รอยสัก-เจาะหู ห้อง {header['ห้อง']}", size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "*" * 21, size=16)

    # สรุปรอยสัก
    tattoo = cfg["summary_tattoo"]
    p = doc.add_paragraph()
    add_text(p, "สรุป\t", bold=True)
    add_text(p, f"{checkbox(tattoo['found']['checked'])} พบรอยสัก ................ คน")
    p2 = doc.add_paragraph()
    add_text(p2, "\t")
    add_text(p2, f"{checkbox(tattoo['none']['checked'])} ไม่มี")

    # สรุปเจาะหู
    piercing = cfg["summary_piercing"]
    p = doc.add_paragraph()
    add_text(p, "สรุป\t", bold=True)
    add_text(p, f"{checkbox(piercing['male_pierced']['checked'])} นักศึกษาชายที่เจาะหู ................ คน")
    p2 = doc.add_paragraph()
    add_text(p2, "\t")
    add_text(p2, f"{checkbox(piercing['female_over_pierced']['checked'])} นักศึกษาหญิงที่เจาะหูเกินข้างละ 1 รู ................ คน")
    p3 = doc.add_paragraph()
    add_text(p3, "\t")
    add_text(p3, f"{checkbox(piercing['none']['checked'])} ไม่มี")

    doc.add_paragraph()

    p = doc.add_paragraph()
    add_text(p, "รายชื่อนักศึกษาที่ตรวจพบรอยสัก", bold=True)

    students = cfg.get("students_with_tattoo", [])
    row_count = max(3, len(students))
    for i in range(row_count):
        name = students[i]["name"] if i < len(students) and students[i].get("name") else ""
        location = students[i]["location"] if i < len(students) and students[i].get("location") else ""
        p = doc.add_paragraph()
        name_part = name if name else "." * 30
        loc_part = location if location else "." * 12
        add_text(p, f"{i + 1}. {name_part}  ตำแหน่งที่พบ {loc_part}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "ลงชื่อ ..................................")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, f"(...................{cfg.get('teacher_name', '')}...................)")


# ---------- สร้างหน้ารูปตามกลุ่ม ----------

def build_group_pages(doc, cfg):
    photos_base = BASE_DIR / cfg.get("photos_dir", "photos")
    sections = cfg.get("photo_sections", [])

    if not sections:
        print("⚠ ไม่พบ photo_sections ใน config.json — ข้ามส่วนรูปภาพ")
        return

    available_w = get_available_width_inches(doc)

    for sec in sections:
        folder_name = sec["folder"]
        group_size = sec["group_size"]
        label_prefix = sec["label_prefix"]

        folder = photos_base / folder_name
        if not folder.exists():
            print(f"⚠ ไม่พบโฟลเดอร์ {folder} — ข้ามส่วนนี้")
            continue

        images = sorted(
            [p for p in folder.iterdir() if p.is_file() and p.suffix.lower() in IMAGE_EXTS],
            key=natural_sort_key,
        )

        if not images:
            print(f"⚠ ไม่มีรูปใน {folder} — ข้ามส่วนนี้")
            continue

        total = len(images)
        if total % group_size != 0:
            last_count = total % group_size
            print(
                f"⚠ {folder_name}: จำนวนรูป ({total}) หารด้วย {group_size} ไม่ลงตัว "
                f"— กลุ่มสุดท้ายมีรูปแค่ {last_count} รูป"
            )

        groups = [images[i : i + group_size] for i in range(0, total, group_size)]

        for gi, group in enumerate(groups, start=1):
            doc.add_page_break()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, f"{label_prefix} {gi}", size=18, bold=True)

            # วางรูป 2 รูปต่อแถว
            for row_start in range(0, len(group), 2):
                row_imgs = group[row_start : row_start + 2]
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_image_row(img_p, row_imgs, available_w)


def main():
    if not CONFIG_PATH.exists():
        sys.exit(f"ไม่พบไฟล์ config.json ที่ {CONFIG_PATH}")

    with open(CONFIG_PATH, encoding="utf-8") as f:
        cfg = json.load(f)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BASE_SIZE)

    build_summary_page(doc, cfg)
    build_group_pages(doc, cfg)

    out_path = BASE_DIR / cfg.get("output_file", "output/รายงาน.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"✅ สร้างรายงานเรียบร้อย: {out_path}")


if __name__ == "__main__":
    main()
